#!/usr/bin/env python3
"""Generate a review docx from normalized issues.json.

Usage:
  python generate_review_doc.py issues.json output.docx [template.docx]

Notes:
- If template.docx is omitted, the script will try to use
  ../assets/docx/report_template.docx relative to this script.
- The template is expected to have:
  1) title / project / date / mapping paragraphs near the top
  2) a table with this header row:
     序号 | 问题描述 | 状态 | 问题截图 | 日期 | 优先级 | 走查人 | 跟进人
- If the template contains a sample issue row, it will be cleared and reused.
"""

from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Cm
from docx.table import _Cell, Table

try:
    from PIL import Image
except ImportError:
    Image = None

HEADERS = ["序号", "问题描述", "状态", "问题截图", "日期", "优先级", "走查人", "跟进人"]
WIDTHS_CM = [1.0, 6.8, 1.8, 7.8, 1.8, 1.4, 2.1, 2.1]
MAX_IMAGE_SIDE_PX = 1600
JPEG_QUALITY = 82


def set_col_width(cell: _Cell, cm: float) -> None:
    cell.width = Cm(cm)


def clear_cell(cell: _Cell) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    for _ in range(len(p.runs)):
        p.runs[0]._element.getparent().remove(p.runs[0]._element)


def resolve_template(custom_path: str | None) -> Path | None:
    if custom_path:
        p = Path(custom_path)
        return p if p.exists() else None
    default_path = Path(__file__).resolve().parent.parent / "assets" / "docx" / "report_template.docx"
    return default_path if default_path.exists() else None


def load_doc(template_path: Path | None) -> Document:
    return Document(str(template_path)) if template_path else Document()


def find_issue_table(doc: Document) -> Table | None:
    for table in doc.tables:
        if len(table.rows) >= 1 and len(table.columns) == len(HEADERS):
            row0 = [c.text.strip() for c in table.rows[0].cells]
            if row0 == HEADERS:
                return table
    return None


def ensure_issue_table(doc: Document) -> Table:
    table = find_issue_table(doc)
    if table is not None:
        return table
    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        table.cell(0, i).text = h
    return table


def fill_header_paragraphs(doc: Document, data: dict) -> None:
    if len(doc.paragraphs) >= 1:
        doc.paragraphs[0].text = data.get("title", "设计走查问题表")
    else:
        doc.add_heading(data.get("title", "设计走查问题表"), level=1)

    if len(doc.paragraphs) >= 2:
        doc.paragraphs[1].text = f'项目：{data.get("project", "")}'
    else:
        doc.add_paragraph(f'项目：{data.get("project", "")}')

    if len(doc.paragraphs) >= 3:
        doc.paragraphs[2].text = f'日期：{data.get("date", "")}'
    else:
        doc.add_paragraph(f'日期：{data.get("date", "")}')

    if len(doc.paragraphs) >= 4:
        doc.paragraphs[3].text = '图片映射：开发页 / 设计稿'
    else:
        doc.add_paragraph('图片映射：开发页 / 设计稿')


def clear_sample_rows(table: Table) -> None:
    while len(table.rows) > 1:
        tbl = table._tbl
        tbl.remove(table.rows[1]._tr)


def prepare_image_for_doc(img_path: Path, temp_dir: Path) -> Path | None:
    if not img_path.exists():
        return None
    if Image is None:
        return img_path

    try:
        with Image.open(img_path) as img:
            img = img.convert("RGB") if img.mode in {"RGBA", "P", "LA"} else img.copy()
            width, height = img.size
            longest = max(width, height)
            if longest > MAX_IMAGE_SIDE_PX:
                scale = MAX_IMAGE_SIDE_PX / float(longest)
                resized = (
                    max(1, int(width * scale)),
                    max(1, int(height * scale)),
                )
                img = img.resize(resized, Image.Resampling.LANCZOS)
            out_path = temp_dir / f"{img_path.stem}_doc.jpg"
            img.save(out_path, format="JPEG", quality=JPEG_QUALITY, optimize=True)
            return out_path
    except Exception:
        return img_path


def add_issue_row(table: Table, row_issue: dict, *, default_date: str, default_reviewer: str, default_owner: str, temp_dir: Path) -> None:
    cells = table.add_row().cells
    for c, w in zip(cells, WIDTHS_CM):
        set_col_width(c, w)

    cells[0].text = str(row_issue.get("seq", ""))
    cells[1].text = str(row_issue.get("description", "")).strip()
    cells[2].text = str(row_issue.get("status", "待修改"))
    cells[4].text = str(row_issue.get("date", default_date))
    cells[5].text = str(row_issue.get("priority", "P1"))
    cells[6].text = str(row_issue.get("reviewer", default_reviewer))
    cells[7].text = str(row_issue.get("owner", default_owner))

    clear_cell(cells[3])
    p = cells[3].paragraphs[0]
    first = True
    for img in row_issue.get("images", [])[:2]:
        img_path = Path(img)
        prepared = prepare_image_for_doc(img_path, temp_dir)
        if prepared is None:
            continue
        if not first:
            p.add_run().add_break()
        run = p.add_run()
        run.add_picture(str(prepared), width=Cm(3.6))
        first = False
    if first:
        cells[3].text = "无截图"


def main() -> int:
    if len(sys.argv) not in {3, 4}:
        print(__doc__)
        return 1

    src = Path(sys.argv[1])
    out = Path(sys.argv[2])
    template_path = resolve_template(sys.argv[3] if len(sys.argv) == 4 else None)

    data = json.loads(src.read_text(encoding="utf-8"))

    doc = load_doc(template_path)
    if not template_path:
        sec = doc.sections[0]
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.5)
        sec.right_margin = Cm(1.5)
        doc.add_heading(data.get("title", "设计走查问题表"), level=1)
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")

    fill_header_paragraphs(doc, data)

    table = ensure_issue_table(doc)
    clear_sample_rows(table)

    default_date = data.get("date", "")
    default_reviewer = data.get("reviewer", "AI 走查助手")
    default_owner = data.get("owner_default", "待指定")

    with tempfile.TemporaryDirectory(prefix="uicheck_doc_") as temp_root:
        temp_dir = Path(temp_root)
        for row_issue in data.get("issues", []):
            add_issue_row(
                table,
                row_issue,
                default_date=default_date,
                default_reviewer=default_reviewer,
                default_owner=default_owner,
                temp_dir=temp_dir,
            )

        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)
    print(f"Saved: {out}")
    if template_path:
        print(f"Template: {template_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
